from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell_border(cell):
    """Apply border to cell"""
    for edge in ['top', 'left', 'bottom', 'right']:
        cell._element.get_or_add_tcPr().append(
            cell._element._new_child(f'w:{edge}')
        )

# Create document
doc = Document()

# Set margins
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Title
title = doc.add_paragraph()
title_run = title.add_run("CRITICAL REVIEW OF RELATED WORKS")
title_run.bold = True
title_run.font.size = Pt(16)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Table headers
headers = [
    "Author & Year",
    "Title of Paper",
    "Problem Addressed",
    "Methodology",
    "Significant Achievement",
    "Limitation",
    "Synthesis"
]

# Create table
table = doc.add_table(rows=6, cols=7)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
col_widths = [1.0, 1.8, 1.5, 1.5, 1.8, 1.5, 1.8]
for i, width in enumerate(col_widths):
    table.columns[i].width = Inches(width)

# Add headers
header_row = table.rows[0]
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
data = [
    [
        "Rashid, Malik & Ghauri (2026)",
        "Ethically integrated generative AI for reading and vocabulary development in higher education",
        "Chronic literacy problems in higher education and lack of empirical evidence on pedagogical effectiveness and ethical use of AI in applied linguistics",
        "Quasi-experimental pretest-posttest control group design with 148 undergraduate students; mixed-methods using TAM framework; 15-week intervention with AI tools",
        "AI group showed significantly higher scores in vocabulary and reading comprehension (p < .001); 30% improvement in learning autonomy; gender differences observed",
        "Short-term intervention may not capture long-term retention; single university context limits generalizability; reliance on self-reported data",
        "Provides empirical evidence for ethically integrated AI tools; supports TAM framework; validates AI effectiveness in skill acquisition"
    ],
    [
        "Ma, Ma, Qi, Zhang & Ruan (2025)",
        "A practical study of AI-based real-time feedback in online physical education teaching",
        "Poor skill acquisition in online physical education due to lack of real-time feedback",
        "8-week randomized controlled trial with 60 students; AI pose recognition system vs. traditional MOOC; mediation analysis using SPSS and Python",
        "AI system significantly enhanced movement quality, fluency, learning interest, and self-directed learning; learning duration identified as primary significant mediator",
        "Technology limitations in accuracy and interactivity; sample limited to undergraduates; short-term intervention; modest sample size (N=60)",
        "Demonstrates AI feedback effectiveness for skill acquisition; mediation analysis provides mechanism understanding; supports AI as cognitive scaffolding tool"
    ],
    [
        "Attwell, Bekiaridis, Deitmer, Perini, Roppertz & Tultys (2020)",
        "Artificial Intelligence in Policies, Processes and Practices of Vocational Education and Training",
        "Limited understanding of AI impact on VET policies, processes, and practices across Europe",
        "Explorative comparative case-study with semi-structured expert interviews in 5 European countries; online survey of VET teachers; literature review",
        "Identified two ways AI enters VET: as teaching tool and as subject of study; developed competency framework for VET teachers; documented good practices",
        "Limited sample size; varying national contexts; AI integration still in early stages; lack of quantitative outcome data",
        "Provides comprehensive overview of AI in VET; identifies teacher competency needs; highlights infrastructure and training challenges"
    ],
    [
        "Setiadi & Yoto (2025)",
        "Building Vocational High School Students Employability Skills in Vocational Education with an Artificial Intelligence (AI) Approach",
        "Need to understand how AI integration impacts employability skills development in vocational high school students",
        "Systematic Literature Review (SLR) of 22 articles from 2019-2025; thematic analysis using PRISMA approach; inclusion/exclusion criteria applied",
        "AI enhances both technical skills (hard skills) through simulations and interactive learning, and interpersonal skills (soft skills) through project-based learning and collaboration",
        "Limited literature specifically on AI in Indonesian vocational context; varying research quality; topic still in infancy",
        "Validates AI effectiveness for employability skills; identifies infrastructure and teacher training as key barriers; supports link and match approach"
    ],
    [
        "PeerLink Synthesis Document (2025)",
        "PeerLink Related Works Synthesis",
        "Need for secure peer-to-peer file sharing with encryption, authentication, and integrity verification",
        "Critical review of 20 studies on P2P file sharing and WebRTC technologies; comparative analysis",
        "Identified that existing systems focus on specific aspects; PeerLink combines WebRTC, E2EE, SHA-256, QR joining, transfer monitoring, and room management",
        "Primarily a synthesis of existing works rather than primary research; requires implementation validation",
        "Provides foundation for secure file transfer in AI learning platforms; combines multiple security features into single practical solution"
    ]
]

# Add data rows
for row_idx, row_data in enumerate(data, start=1):
    row = table.rows[row_idx]
    for col_idx, text in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Research Gap section
gap_title = doc.add_paragraph()
gap_run = gap_title.add_run("RESEARCH GAP / OVERALL SYNTHESIS")
gap_run.bold = True
gap_run.font.size = Pt(14)

doc.add_paragraph()

gap_text = """The reviewed works demonstrate that AI-based learning systems significantly enhance skill acquisition in vocational and higher education contexts. Rashid et al. (2026) provide empirical evidence for ethically integrated AI tools in improving vocabulary and reading comprehension, while Ma et al. (2025) demonstrate the effectiveness of AI feedback systems for skill acquisition through increased learning duration. Attwell et al. (2020) establish the policy and competency frameworks necessary for AI integration in VET, and Setiadi & Yoto (2025) validate AI's role in developing both technical and interpersonal employability skills.

However, existing studies primarily focus on general vocational education, language learning, and physical education, with limited attention to AI-based learning systems specifically designed for cream, soap, and perfume production skills. Additionally, while individual AI features such as feedback, prediction, and recommendations have been studied, there is a lack of integrated platforms that combine multiple AI functionalities (skill prediction, lesson recommendation, progress tracking, and ethical considerations) into a comprehensive learning system for practical skill acquisition.

There is therefore a gap in research on AI-based learning systems that integrate prediction, recommendation, progress tracking, and ethical considerations for teaching cream, soap, and perfume production skills. This study aims to address this gap by designing and implementing such a system."""

doc.add_paragraph(gap_text)

# Save
doc.save("Critical_Review_of_Related_Works.docx")
print("✅ Word document created: Critical_Review_of_Related_Works.docx")